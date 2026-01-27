#!/usr/bin/env python3
"""
Resume Parser - Extracts structured information from PDF and DOCX resumes
Uses pdfplumber for PDFs and python-docx for Word documents
"""

import sys
import json
import re
import os
from pathlib import Path

# PDF parsing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# DOCX parsing
try:
    from docx import Document
except ImportError:
    Document = None


class ResumeParser:
    def __init__(self):
        # Common skills to look for
        self.common_skills = [
            # Programming Languages
            'JavaScript', 'Python', 'Java', 'C++', 'C#', 'Ruby', 'PHP', 'Swift',
            'Kotlin', 'Go', 'Rust', 'TypeScript', 'Scala', 'R', 'MATLAB', 'Perl',

            # Web Technologies
            'React', 'Angular', 'Vue', 'Vue.js', 'Node.js', 'Express', 'Django',
            'Flask', 'Spring', 'Laravel', 'Rails', 'ASP.NET', 'Next.js', 'Nuxt.js',

            # Frontend
            'HTML', 'CSS', 'SASS', 'SCSS', 'Less', 'Bootstrap', 'Tailwind',
            'jQuery', 'Redux', 'GraphQL', 'REST', 'RESTful',

            # Databases
            'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle',
            'SQLite', 'Cassandra', 'DynamoDB', 'Firebase', 'Elasticsearch',

            # Cloud & DevOps
            'AWS', 'Azure', 'GCP', 'Google Cloud', 'Docker', 'Kubernetes',
            'Jenkins', 'CI/CD', 'Terraform', 'Ansible', 'Linux', 'Unix',
            'Git', 'GitHub', 'GitLab', 'Bitbucket',

            # Data Science & ML
            'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch',
            'Keras', 'Scikit-learn', 'Pandas', 'NumPy', 'Data Analysis',
            'Data Science', 'NLP', 'Computer Vision', 'AI', 'Statistics',

            # Mobile
            'iOS', 'Android', 'React Native', 'Flutter', 'Xamarin',

            # Tools & Others
            'Jira', 'Confluence', 'Trello', 'Slack', 'Figma', 'Photoshop',
            'Illustrator', 'Sketch', 'Adobe XD', 'InVision',

            # Methodologies
            'Agile', 'Scrum', 'Kanban', 'Waterfall', 'TDD', 'BDD',

            # Soft Skills
            'Leadership', 'Communication', 'Team Management', 'Problem Solving',
            'Project Management', 'Critical Thinking', 'Analytical Skills',
        ]

        # Education keywords
        self.education_levels = {
            'phd': ['ph.d', 'phd', 'doctorate', 'doctoral'],
            'master': ['master', 'm.s.', 'm.s', 'msc', 'mba', 'm.a.', 'ma'],
            'bachelor': ['bachelor', 'b.s.', 'b.s', 'bsc', 'b.a.', 'ba', 'b.tech', 'btech'],
            'associate': ['associate', 'a.s.', 'a.a.'],
            'high-school': ['high school', 'secondary', 'diploma', 'ged'],
        }

        # Job categories mapping
        self.category_keywords = {
            'Software Development': ['software', 'developer', 'programming', 'engineer', 'code', 'coding', 'full stack', 'frontend', 'backend', 'web developer'],
            'Data Science': ['data science', 'data scientist', 'machine learning', 'ml', 'ai', 'artificial intelligence', 'analytics', 'data analysis', 'big data'],
            'Design': ['design', 'designer', 'ui', 'ux', 'user experience', 'user interface', 'graphic', 'creative', 'visual'],
            'Marketing': ['marketing', 'seo', 'sem', 'content', 'social media', 'digital marketing', 'brand', 'advertising'],
            'Sales': ['sales', 'business development', 'account manager', 'client', 'revenue', 'crm'],
            'Finance': ['finance', 'financial', 'accounting', 'accountant', 'auditor', 'tax', 'banking', 'investment'],
            'HR': ['human resources', 'hr', 'recruitment', 'recruiter', 'talent', 'hiring', 'people operations'],
            'Operations': ['operations', 'logistics', 'supply chain', 'inventory', 'procurement'],
            'Customer Service': ['customer service', 'customer support', 'helpdesk', 'client support', 'customer success'],
            'Healthcare': ['healthcare', 'medical', 'nursing', 'nurse', 'doctor', 'physician', 'clinical', 'pharmacy'],
            'Education': ['education', 'teacher', 'professor', 'instructor', 'training', 'curriculum', 'academic'],
            'Engineering': ['engineering', 'mechanical', 'electrical', 'civil', 'chemical', 'structural', 'cad'],
        }

        # Experience level patterns
        self.experience_patterns = [
            (r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s+(?:experience|exp)', 'years'),
            (r'(?:experience|exp)[\s:]+(\d+)\+?\s*(?:years?|yrs?)', 'years'),
            (r'(\d+)\+?\s*(?:years?|yrs?)(?:\s+in)?', 'years'),
        ]

    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file using pdfplumber"""
        if pdfplumber is None:
            raise ImportError("pdfplumber is not installed")

        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text

    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file using python-docx"""
        if Document is None:
            raise ImportError("python-docx is not installed")

        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        return text

    def extract_text(self, file_path):
        """Extract text based on file extension"""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def extract_skills(self, text):
        """Extract skills from resume text"""
        found_skills = []
        text_lower = text.lower()

        for skill in self.common_skills:
            # Create pattern that matches whole word
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                found_skills.append(skill)

        # Remove duplicates while preserving order
        seen = set()
        unique_skills = []
        for skill in found_skills:
            if skill.lower() not in seen:
                seen.add(skill.lower())
                unique_skills.append(skill)

        return unique_skills[:15]  # Return top 15 skills

    def extract_job_title(self, text):
        """Extract current/recent job title"""
        lines = text.split('\n')

        # Common job title patterns
        title_patterns = [
            r'(?:current|present|latest)?\s*(?:position|title|role|designation)[\s:]+(.+)',
            r'^([\w\s]+(?:engineer|developer|manager|analyst|designer|director|specialist|consultant|architect|lead|senior|junior|associate|coordinator|administrator|executive))',
        ]

        for line in lines[:30]:  # Check first 30 lines
            line = line.strip()
            if not line:
                continue

            for pattern in title_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    title = match.group(1).strip()
                    # Clean up the title
                    title = re.sub(r'\s+', ' ', title)
                    if len(title) > 5 and len(title) < 100:
                        return title

        return ""

    def extract_company(self, text):
        """Extract current/recent company name"""
        patterns = [
            r'(?:current|present)?\s*(?:company|employer|organization)[\s:]+(.+)',
            r'(?:working|worked)\s+(?:at|for|with)\s+([A-Z][\w\s&.,]+?)(?:\s+as|\s+from|\s+since|\.|\n)',
            r'(?:employed|employment)\s+(?:at|with)\s+([A-Z][\w\s&.,]+?)(?:\s+as|\s+from|\.|,|\n)',
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                company = match.group(1).strip()
                # Clean up
                company = re.sub(r'\s+', ' ', company)
                company = company.rstrip('.,')
                if len(company) > 2 and len(company) < 100:
                    return company

        return ""

    def extract_experience_years(self, text):
        """Extract years of experience and return experience level"""
        for pattern, _ in self.experience_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    years = int(match.group(1))
                    if years <= 2:
                        return "Entry Level (0-2 years)"
                    elif years <= 5:
                        return "Mid Level (3-5 years)"
                    elif years <= 10:
                        return "Senior Level (6-10 years)"
                    else:
                        return "Executive Level (10+ years)"
                except (ValueError, IndexError):
                    pass

        return ""

    def extract_education(self, text):
        """Extract education level"""
        text_lower = text.lower()

        # Check for each education level
        for level, keywords in self.education_levels.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return level

        return ""

    def extract_university(self, text):
        """Extract university/institution name"""
        patterns = [
            r'(?:university|college|institute|school)\s+(?:of\s+)?([A-Z][\w\s,]+)',
            r'([A-Z][\w\s]+(?:University|College|Institute|School))',
            r'(?:graduated|studied|degree)\s+(?:from|at)\s+([A-Z][\w\s,]+)',
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.MULTILINE)
            if match:
                university = match.group(1).strip() if match.group(1) else match.group(0).strip()
                # Clean up
                university = re.sub(r'\s+', ' ', university)
                university = university.rstrip('.,')
                if len(university) > 3 and len(university) < 150:
                    return university

        return ""

    def extract_graduation_year(self, text):
        """Extract graduation year"""
        # Look for years in education context
        patterns = [
            r'(?:graduated?|class\s+of|batch\s+of|graduation)[\s:]+(\d{4})',
            r'(\d{4})[\s-]+(?:present|current|now)',
            r'(\d{4})\s*[-–]\s*(\d{4})',
        ]

        years = []
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple):
                    for m in match:
                        if m and m.isdigit():
                            year = int(m)
                            if 1970 <= year <= 2030:
                                years.append(year)
                else:
                    if match.isdigit():
                        year = int(match)
                        if 1970 <= year <= 2030:
                            years.append(year)

        # Also look for standalone 4-digit years
        standalone_years = re.findall(r'\b(19\d{2}|20\d{2})\b', text)
        for y in standalone_years:
            year = int(y)
            if 1970 <= year <= 2030:
                years.append(year)

        if years:
            # Return the most recent year that's not in the future
            import datetime
            current_year = datetime.datetime.now().year
            valid_years = [y for y in years if y <= current_year + 4]
            if valid_years:
                return str(max(valid_years))

        return ""

    def extract_job_categories(self, text, skills):
        """Infer job categories from text and skills"""
        categories = []
        text_lower = text.lower()
        skills_lower = [s.lower() for s in skills]

        for category, keywords in self.category_keywords.items():
            for keyword in keywords:
                if keyword in text_lower or keyword in ' '.join(skills_lower):
                    if category not in categories:
                        categories.append(category)
                        break

        return categories if categories else ['Other']

    def parse(self, file_path):
        """Main parsing function"""
        try:
            # Extract text from file
            text = self.extract_text(file_path)

            if not text or len(text.strip()) < 50:
                return {
                    'success': False,
                    'error': 'Could not extract sufficient text from the file'
                }

            # Extract all components
            skills = self.extract_skills(text)
            job_title = self.extract_job_title(text)
            company = self.extract_company(text)
            experience = self.extract_experience_years(text)
            education = self.extract_education(text)
            university = self.extract_university(text)
            graduation_year = self.extract_graduation_year(text)
            job_categories = self.extract_job_categories(text, skills)

            return {
                'success': True,
                'data': {
                    'currentJobTitle': job_title,
                    'experience': experience,
                    'currentCompany': company,
                    'expectedSalary': '',  # Not typically in resumes
                    'skills': skills,
                    'education': education,
                    'university': university,
                    'graduationYear': graduation_year,
                    'jobCategories': job_categories,
                    'bio': ''
                }
            }

        except ImportError as e:
            return {
                'success': False,
                'error': f'Missing required library: {str(e)}'
            }
        except Exception as e:
            return {
                'success': False,
                'error': f'Error parsing resume: {str(e)}'
            }


def main():
    if len(sys.argv) < 2:
        print(json.dumps({
            'success': False,
            'error': 'No file path provided'
        }))
        sys.exit(1)

    file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(json.dumps({
            'success': False,
            'error': f'File not found: {file_path}'
        }))
        sys.exit(1)

    parser = ResumeParser()
    result = parser.parse(file_path)

    print(json.dumps(result))


if __name__ == '__main__':
    main()
